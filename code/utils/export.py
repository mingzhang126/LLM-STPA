from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import json
import csv
import os
import logging

logger = logging.getLogger(__name__)

# 注册中文字体（msyh，微软雅黑）
try:
    font_paths = [
        "C:\\Windows\\Fonts\\msyh.ttc",  # Windows
        "/home/zhangming/code/stpa_html/utils/MSYH.TTC",  # Linux
    ]
    font_registered = False
    for font_path in font_paths:
        if os.path.exists(font_path):
            pdfmetrics.registerFont(TTFont("msyh", font_path))
            logger.debug(f"成功注册 微软雅黑 字体: {font_path}")
            font_registered = True
            break
    if not font_registered:
        raise Exception("未找到 微软雅黑 或 Noto Sans CJK 字体文件")
except Exception as e:
    logger.error(f"注册字体失败: {str(e)}")
    raise Exception("无法加载中文字体，请确保字体文件存在")

def export_to_pdf(results, system_description, lang, filename):
    doc = SimpleDocTemplate(filename, pagesize=A4)
    styles = getSampleStyleSheet()
    chinese_style = ParagraphStyle(
        name="Chinese",
        parent=styles["Normal"],
        fontName="msyh",
        fontSize=10,  # 减小字体以适应宽表格
        leading=12,
        spaceAfter=6
    )
    chinese_heading = ParagraphStyle(
        name="ChineseHeading",
        parent=styles["Heading2"],
        fontName="msyh",
        fontSize=12,
        leading=14,
        spaceAfter=8
    )
    chinese_title = ParagraphStyle(
        name="ChineseTitle",
        parent=styles["Title"],
        fontName="msyh",
        fontSize=16,
        leading=18,
        spaceAfter=12
    )

    story = []
    story.append(Paragraph(lang["title"], chinese_title))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"{lang['system_desc']}: {system_description}", chinese_style))
    story.append(Spacer(1, 12))

    for step in [1, 2, 3, 4, 5, 6]:
        step_key = str(step)
        step_name = lang[f"step{step}"]
        story.append(Paragraph(f"步骤 {step}: {step_name}", chinese_heading))
        story.append(Paragraph(results.get(step_key, "未生成"), chinese_style))
        story.append(Spacer(1, 12))

       
        if step == 1:
             # 步骤 1：Losses 表格
            losses_path = results.get(f"{step}_losses_path")
            if losses_path and os.path.exists(losses_path):
                with open(losses_path, 'r', encoding='utf-8') as f:
                    losses = json.load(f)
            else:
                losses = results.get(f"{step}_losses", [])
            if losses:
                logger.debug(f"渲染步骤 1 Losses 表格: {losses}")
                story.append(Paragraph(lang["losses_title"], chinese_heading))
                data = [[lang["loss_id"], lang["description"]]]
                for row in losses:
                    data.append([row["id"], row["description"]])
                table = Table(data, colWidths=[100, 400])  # 调整列宽
                table.setStyle(TableStyle([
                    ("GRID", (0, 0), (-1, -1), 1, colors.black),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                    ("FONTNAME", (0, 0), (-1, -1), "msyh"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("BACKGROUND", (0, 1), (-1, -1), colors.lightgrey),
                    ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("BOX", (0, 0), (-1, -1), 1, colors.black),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4)
                ]))
                story.append(table)
                story.append(Spacer(1, 12))
            # 步骤 1：Hazards 表格
            hazards_path = results.get(f"{step}_hazards_path")
            if hazards_path and os.path.exists(hazards_path):
                with open(hazards_path, 'r', encoding='utf-8') as f:
                    hazards = json.load(f)
            else:
                hazards = results.get(f"{step}_hazards", [])
            if hazards:
                logger.debug(f"渲染步骤 1 Hazards 表格: {hazards}")
                story.append(Paragraph(lang["hazards_title"], chinese_heading))
                data = [[lang["hazard_id"], lang["description"], lang["linked_losses"]]]
                for row in hazards:
                    data.append([row["id"], row["description"], ", ".join(row["linked_losses"])])    
                table = Table(data, colWidths=[100, 300, 100])  # 调整列宽
                table.setStyle(TableStyle([
                    ("GRID", (0, 0), (-1, -1), 1, colors.black),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                    ("FONTNAME", (0, 0), (-1, -1), "msyh"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("BACKGROUND", (0, 1), (-1, -1), colors.lightgrey),
                    ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("BOX", (0, 0), (-1, -1), 1, colors.black),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4)
                ]))
                story.append(table)
                story.append(Spacer(1, 12))
            # 步骤 1：Safety Constraints 表格
            constraints_path = results.get(f"{step}_safety_constraints_path")
            if constraints_path and os.path.exists(constraints_path):
                with open(constraints_path, 'r', encoding='utf-8') as f:
                    safety_constraints = json.load(f)
            else:
                safety_constraints = results.get(f"{step}_safety_constraints", [])
            if safety_constraints:
                logger.debug(f"渲染步骤 1 Safety Constraints 表格: {safety_constraints}")
                story.append(Paragraph(lang["safety_constraints_title"], chinese_heading))
                data = [[lang["loss_id"], lang["description"]]]
                for row in safety_constraints:
                    data.append([row["id"], row["description"]])
                table = Table(data, colWidths=[100, 400])  # 调整列宽
                table.setStyle(TableStyle([
                    ("GRID", (0, 0), (-1, -1), 1, colors.black),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                    ("FONTNAME", (0, 0), (-1, -1), "msyh"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("BACKGROUND", (0, 1), (-1, -1), colors.lightgrey),
                    ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("BOX", (0, 0), (-1, -1), 1, colors.black),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4)
                ]))
                story.append(table)
                story.append(Spacer(1, 12))
        # 步骤 2：控制结构表格
        if step == 2:
            table_path = results.get(f"{step}_table_path")
            if table_path and os.path.exists(table_path):
                with open(table_path, 'r', encoding='utf-8') as f:
                    table_data = json.load(f)
            else:
                table_data = results.get(f"{step}_table", [])
            if table_data:
                logger.debug(f"渲染步骤 2 表格: {table_data}")
                story.append(Paragraph(lang["step2"], chinese_heading))
                data = [
                    [
                        lang["table_id"],
                        lang["table_control_action"],
                        lang["table_from"],
                        lang["table_to"],
                        lang["table_condition"]
                    ]
                ]
                for row in table_data:
                    data.append([
                        row["id"],
                        row["control_action"],
                        row["from"],
                        row["to"],
                        row["condition"]
                    ])
                table = Table(data, colWidths=[80, 120, 100, 100, 100])  # 调整列宽
                table.setStyle(TableStyle([
                    ("GRID", (0, 0), (-1, -1), 1, colors.black),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                    ("FONTNAME", (0, 0), (-1, -1), "msyh"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("BACKGROUND", (0, 1), (-1, -1), colors.lightgrey),
                    ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("BOX", (0, 0), (-1, -1), 1, colors.black),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4)
                ]))
                story.append(table)
                story.append(Spacer(1, 12))

        # 步骤 3：不安全控制行为表格
        if step == 3:
            table_path = results.get(f"{step}_table_path")
            if table_path and os.path.exists(table_path):
                with open(table_path, 'r', encoding='utf-8') as f:
                    table_data = json.load(f)
            else:
                table_data = results.get(f"{step}_table", [])
            logger.debug(f"渲染步骤 3 表格: {table_data}")
            story.append(Paragraph(lang["step3"], chinese_heading))
            data = [
                [
                    lang["table_id"],
                    lang["table_control_action"],
                    lang["table_cant_providing"],
                    lang["table_providing"],
                    lang["table_wrong_moment"],
                    lang["table_sustained_abnormal"]
                ]
            ]
            for row in table_data:
                data.append([
                    row["id"],
                    row["control_action"],
                    row["cant_providing"],
                    row["providing"],
                    row["wrong_moment"],
                    row["sustained_abnormal"]
                ])
            table = Table(data, colWidths=[80, 100, 100, 100, 100, 100])  # 调整列宽
            table.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
                ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("FONTNAME", (0, 0), (-1, -1), "msyh"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("BACKGROUND", (0, 1), (-1, -1), colors.lightgrey),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("BOX", (0, 0), (-1, -1), 1, colors.black),
                ("LEFTPADDING", (0, 0), (-1, -1), 3),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ("WORDWRAP", (0, 0), (-1, -1), True)  # 启用自动换行
            ]))
            story.append(table)
            story.append(Spacer(1, 12))

        # 步骤 4：HARA 评估表格
        if step == 4:
            table_path = results.get(f"{step}_table_path")
            if table_path and os.path.exists(table_path):
                with open(table_path, 'r', encoding='utf-8') as f:
                    table_data = json.load(f)
            else:
                table_data = results.get(f"{step}_table", [])
            logger.debug(f"渲染步骤 4 表格: {table_data}")
            story.append(Paragraph(lang["step4"], chinese_heading))
            data = [
                [
                    lang["table_id"],
                    lang["table_control_action"],
                    lang["table_uca"],
                    lang["table_hazard"],
                    lang["table_scenario"],
                    lang["table_impact"],
                    lang["table_severity"],
                    lang["table_severity_desc"],
                    lang["table_controllability"],
                    lang["table_controllability_desc"],
                    lang["table_accepted"]
                ]
            ]
            for row in table_data:
                data.append([
                    row["id"],
                    row["control_action"],
                    row["uca"],
                    row["hazard"],
                    row["scenario"],
                    row["impact"],
                    row["severity"],
                    row["severity_desc"],
                    row["controllability"],
                    row["controllability_desc"],
                    row["accepted"]
                ])
            table = Table(data, colWidths=[50, 60, 60, 60, 60, 60, 50, 60, 50, 60, 50])  # 调整列宽以适应 11 列
            table.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
                ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("FONTNAME", (0, 0), (-1, -1), "msyh"),
                ("FONTSIZE", (0, 0), (-1, -1), 7),  # 减小字体以适应宽表格
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("BACKGROUND", (0, 1), (-1, -1), colors.lightgrey),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("BOX", (0, 0), (-1, -1), 1, colors.black),
                ("LEFTPADDING", (0, 0), (-1, -1), 3),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ("WORDWRAP", (0, 0), (-1, -1), True)  # 启用自动换行
            ]))
            story.append(table)
            story.append(Spacer(1, 12))

        # 步骤 5：致因因素表格
        if step == 5:
            table_path = results.get(f"{step}_table_path")
            if table_path and os.path.exists(table_path):
                with open(table_path, 'r', encoding='utf-8') as f:
                    table_data = json.load(f)
            else:
                table_data = results.get(f"{step}_table", [])
            logger.debug(f"渲染步骤 5 表格: {table_data}")
            story.append(Paragraph(lang["step5"], chinese_heading))
            data = [[
                lang["table_id"],
                lang["table_misuse_scenario"],
                lang["table_impact"],
                lang["table_stakeholder"],
                lang["table_is_analyzed"],
                lang["table_analysis_reason"],
                lang["table_defect_or_misuse"],
                lang["table_misuse_process"],
                lang["table_guide_word"],
                lang["table_causal_factor"],
                lang["table_requirement"]
            ]]
            for row in table_data:
                data.append([
                    row["id"],
                    row["misuse_scenario"],
                    row["impact"],
                    row["stakeholder"],
                    row["is_analyzed"],
                    row["analysis_reason"],
                    row["defect_or_misuse"],
                    row["misuse_process"],
                    row["guide_word"],
                    row["causal_factor"],
                    row["requirement"]
                ])
            table = Table(data, colWidths=[50, 60, 60, 60, 50, 60, 60, 60, 60, 60, 60])  # 调整列宽以适应 11 列
            table.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
                ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("FONTNAME", (0, 0), (-1, -1), "msyh"),
                ("FONTSIZE", (0, 0), (-1, -1), 7),  # 减小字体以适应宽表格
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("BACKGROUND", (0, 1), (-1, -1), colors.lightgrey),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("BOX", (0, 0), (-1, -1), 1, colors.black),
                ("LEFTPADDING", (0, 0), (-1, -1), 3),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ("WORDWRAP", (0, 0), (-1, -1), True)  # 启用自动换行
            ]))
            story.append(table)
            story.append(Spacer(1, 12))

        # 步骤 6：安全需求表格
        if step == 6:
            table_path = results.get(f"{step}_table_path")
            if table_path and os.path.exists(table_path):
                with open(table_path, 'r', encoding='utf-8') as f:
                    table_data = json.load(f)
            else:
                table_data = results.get(f"{step}_table", [])
            logger.debug(f"渲染步骤 6 表格: {table_data}")
            story.append(Paragraph(lang["step6"], chinese_heading))
            data = [[lang["table_id"], lang["table_safety_req"], lang["table_tc_id"], lang["table_trigger"]]]
            for row in table_data:
                data.append([row["id"], row["safety_req"], row["tc_id"], row["trigger_condition"]])
            table = Table(data, colWidths=[100, 150, 100, 150])  # 调整列宽
            table.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
                ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("FONTNAME", (0, 0), (-1, -1), "msyh"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("BACKGROUND", (0, 1), (-1, -1), colors.lightgrey),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("BOX", (0, 0), (-1, -1), 1, colors.black),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4)
            ]))
            story.append(table)
            story.append(Spacer(1, 12))
    logger.debug(f"开始生成 PDF: {filename}")
    try:
        doc.build(story)
        logger.debug(f"PDF 生成完成: {filename}")
    except Exception as e:
        logger.error(f"PDF 生成失败: {str(e)}")
        raise

def export_to_word(results, system_description, lang, filename):
    doc = Document()
    # 设置文档默认字体为 微软雅黑（中文） 和 Times New Roman（英文）
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    r = style.element.rPr
    rFonts = r.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), '微软雅黑')  # 设置中文字体
    doc.add_heading(lang["title"], 0)
    doc.add_paragraph(f"{lang['system_desc']}: {system_description}")

    for step in [1, 2, 3, 4, 5, 6]:
        step_key = str(step)
        step_name = lang[f"step{step}"]
        doc.add_heading(f"步骤 {step}: {step_name}", level=2)
        doc.add_paragraph(results.get(step_key, "未生成"))

        if step == 1:
            for sub_step in ["losses", "hazards", "safety_constraints"]:
                sub_path = results.get(f"{step}_{sub_step}_path")
                if sub_path and os.path.exists(sub_path):
                    with open(sub_path, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                    doc.add_heading(lang[f"{sub_step}_title"], level=3)
                    table = doc.add_table(rows=1, cols=3 if sub_step == "hazards" else 2)
                    table.style = "Table Grid"
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = lang["loss_id" if sub_step != "hazards" else "hazard_id"]
                    hdr_cells[1].text = lang["description"]
                    if sub_step == "hazards":
                        hdr_cells[2].text = lang["linked_losses"]
                    for row in data:
                        row_cells = table.add_row().cells
                        row_cells[0].text = row.get("id", "")
                        row_cells[1].text = row.get("description", "")
                        if sub_step == "hazards":
                            row_cells[2].text = ", ".join(row.get("linked_losses", []))
        
        table_path = results.get(f"{step}_table_path")
        if step > 1 and table_path and os.path.exists(table_path):
            with open(table_path, 'r', encoding='utf-8') as f:
                table_data = json.load(f)
            doc.add_heading(lang[f"step{step}"], level=3)
            cols = {2: 5, 3: 6, 4: 11, 5: 11, 6: 4}[step]
            table = doc.add_table(rows=1, cols=cols)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            if step == 2:
                headers = ["table_id", "table_control_action", "table_from", "table_to", "table_condition"]
            elif step == 3:
                headers = ["table_id", "table_control_action", "table_cant_providing", "table_providing", "table_wrong_moment", "table_sustained_abnormal"]
            elif step == 4:
                headers = ["table_id", "table_control_action", "table_uca", "table_hazard", "table_scenario", "table_impact", "table_severity", "table_severity_desc", "table_controllability", "table_controllability_desc", "table_accepted"]
            elif step == 5:
                headers = ["table_id", "table_misuse_scenario", "table_impact", "table_stakeholder", "table_is_analyzed", "table_analysis_reason", "table_defect_or_misuse", "table_misuse_process", "table_guide_word", "table_causal_factor", "table_requirement"]
            elif step == 6:
                headers = ["table_id", "table_safety_req", "table_tc_id", "table_trigger_condition"]
            for i, key in enumerate(headers):
                hdr_cells[i].text = lang[key]
            for row in table_data:
                row_cells = table.add_row().cells
                for i, key in enumerate([h.replace("table_", "") for h in headers]):
                    row_cells[i].text = row.get(key, "")

    doc.save(filename)
    logger.debug(f"生成 Word: {filename}")

def export_to_json(results, system_description, lang, filename):
    for step in [1, 2, 3, 4, 5, 6]:
        table_path = results.get(f"{step}_table")
        if table_path and os.path.exists(table_path):
            with open(table_path, 'r', encoding='utf-8') as f:
                table_data = json.load(f)
        else:
            table_data = results.get(f"{step}_table", [])
        data = {
            lang["system_desc"]: system_description,
            "steps": {
                lang[f"step{step}"]: {
                    "text": results.get(str(step), "未生成"),
                    "losses": results.get(f"{step}_losses") if step == 1 else None,
                    "hazards": results.get(f"{step}_hazards") if step == 1 else None,
                    "safety_constraints": results.get(f"{step}_safety_constraints") if step == 1 else None,
                    "table": results.get(f"{step}_table"),
                    "hara": results.get(f"{step}_hara") if step == 4 else None
                }
            }
        }
    with open(filename, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    logger.debug(f"生成 JSON: {filename}")

def export_to_csv(results, system_description, lang, filename):
    with open(filename, "w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f)
        writer.writerow([lang["system_desc"], system_description])
        writer.writerow([])

        for step in [1, 2, 3, 4, 5, 6]:
            step_key = str(step)
            writer.writerow([f"步骤 {step}: {lang[f'step{step}']}"])
            writer.writerow([results.get(step_key, "未生成")])
            writer.writerow([])

            # 步骤 1：Losses 表格
            if step == 1:
                table_path = results.get(f"{step}_losses")
                if table_path and os.path.exists(table_path):
                    with open(table_path, 'r', encoding='utf-8') as f:
                        table_data = json.load(f)
                else:
                    table_data = results.get(f"{step}_losses", [])
                
                writer.writerow([lang["losses_title"]])
                writer.writerow([lang["loss_id"], lang["description"]])
                for row in table_data:
                    writer.writerow([row["id"], row["description"]])
                writer.writerow([])

            # 步骤 1：Hazards 表格
            if step == 1:
                table_path = results.get(f"{step}_hazards")
                if table_path and os.path.exists(table_path):
                    with open(table_path, 'r', encoding='utf-8') as f:
                        table_data = json.load(f)
                else:
                    table_data = results.get(f"{step}_hazards", [])
                writer.writerow([lang["hazards_title"]])
                writer.writerow([lang["hazard_id"], lang["description"], lang["linked_losses"]])
                for row in table_data:
                    writer.writerow([row["id"], row["description"], ", ".join(row["linked_losses"])])
                writer.writerow([])

            # 步骤 1：Safety Constraints 表格
            if step == 1:
                table_path = results.get(f"{step}_safety_constraints")
                if table_path and os.path.exists(table_path):
                    with open(table_path, 'r', encoding='utf-8') as f:
                        table_data = json.load(f)
                else:
                    table_data = results.get(f"{step}_safety_constraints", [])
                writer.writerow([lang["safety_constraints_title"]])
                writer.writerow([lang["loss_id"], lang["description"]])
                for row in table_data:
                    writer.writerow([row["id"], row["description"]])
                writer.writerow([])

            # 步骤 2：控制结构表格
            if step == 2:
                table_path = results.get(f"{step}_table")
                if table_path and os.path.exists(table_path):
                    with open(table_path, 'r', encoding='utf-8') as f:
                        table_data = json.load(f)
                else:
                    table_data = results.get(f"{step}_table", [])
                writer.writerow([
                    lang["table_id"],
                    lang["table_control_action"],
                    lang["table_from"],
                    lang["table_to"],
                    lang["table_condition"]
                ])
                for row in table_data:
                    writer.writerow([
                        row["id"],
                        row["control_action"],
                        row["from"],
                        row["to"],
                        row["condition"]
                    ])
                writer.writerow([])

            # 步骤 3：不安全控制行为表格
            if step == 3 and results.get(f"{step}_table"):
                table_path = results.get(f"{step}_table")
                if table_path and os.path.exists(table_path):
                    with open(table_path, 'r', encoding='utf-8') as f:
                        table_data = json.load(f)
                else:
                    table_data = results.get(f"{step}_table", [])
                writer.writerow([
                    lang["table_id"],
                    lang["table_control_action"],
                    lang["table_cant_providing"],
                    lang["table_providing"],
                    lang["table_wrong_moment"],
                    lang["table_sustained_abnormal"]
                ])
                for row in table_data:
                    writer.writerow([
                        row["id"],
                        row["control_action"],
                        row["cant_providing"],
                        row["providing"],
                        row["wrong_moment"],
                        row["sustained_abnormal"]
                    ])
                writer.writerow([])

            # 步骤 4：HARA 评估表格
            if step == 4 and results.get(f"{step}_table"):
                table_path = results.get(f"{step}_table")
                if table_path and os.path.exists(table_path):
                    with open(table_path, 'r', encoding='utf-8') as f:
                        table_data = json.load(f)
                else:
                    table_data = results.get(f"{step}_table", [])
                writer.writerow([
                    lang["table_id"],
                    lang["table_control_action"],
                    lang["table_uca"],
                    lang["table_hazard"],
                    lang["table_scenario"],
                    lang["table_impact"],
                    lang["table_severity"],
                    lang["table_severity_desc"],
                    lang["table_controllability"],
                    lang["table_controllability_desc"],
                    lang["table_accepted"]
                ])
                for row in table_data:
                    writer.writerow([
                        row["id"],
                        row["control_action"],
                        row["uca"],
                        row["hazard"],
                        row["scenario"],
                        row["impact"],
                        row["severity"],
                        row["severity_desc"],
                        row["controllability"],
                        row["controllability_desc"],
                        row["accepted"]
                    ])
                writer.writerow([])

            # 步骤 5：致因因素表格
            if step == 5 and results.get(f"{step}_table"):
                table_path = results.get(f"{step}_table")
                if table_path and os.path.exists(table_path):
                    with open(table_path, 'r', encoding='utf-8') as f:
                        table_data = json.load(f)
                else:
                    table_data = results.get(f"{step}_table", [])
                writer.writerow([
                    lang["table_id"],
                    lang["table_misuse_scenario"],
                    lang["table_impact"],
                    lang["table_stakeholder"],
                    lang["table_is_analyzed"],
                    lang["table_analysis_reason"],
                    lang["table_defect_or_misuse"],
                    lang["table_misuse_process"],
                    lang["table_guide_word"],
                    lang["table_causal_factor"],
                    lang["table_requirement"]
                ])
                for row in table_data:
                    writer.writerow([
                        row["id"],
                        row["misuse_scenario"],
                        row["impact"],
                        row["stakeholder"],
                        row["is_analyzed"],
                        row["analysis_reason"],
                        row["defect_or_misuse"],
                        row["misuse_process"],
                        row["guide_word"],
                        row["causal_factor"],
                        row["requirement"]
                    ])
                writer.writerow([])

            # 步骤 6：安全需求表格
            if step == 6 and results.get(f"{step}_table"):
                table_path = results.get(f"{step}_table")
                if table_path and os.path.exists(table_path):
                    with open(table_path, 'r', encoding='utf-8') as f:
                        table_data = json.load(f)
                else:
                    table_data = results.get(f"{step}_table", [])
                writer.writerow([lang["table_id"], lang["table_safety_req"], lang["table_tc_id"], lang["table_trigger"]])
                for row in table_data:
                    writer.writerow([row["id"], row["safety_req"], row["tc_id"], row["trigger_condition"]])
                writer.writerow([])

    logger.debug(f"生成 CSV: {filename}")